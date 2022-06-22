import pandas as pd
from docx import Document
from docx.shared import Inches
import datetime as dt

if __name__ == '__main__':

    FONT_SIZE = 5

    # Read Excel in pandas dataframe
    df = pd.read_excel('./Originals/input_database.xlsx')

    # Remove NAN values (not a number)
    df.fillna('', inplace=True)

    # Instantiate a document object
    letter = Document()

    # Each row is a tuple of (index, pandas sries)
    for idx, entry in df.iterrows():

        # We go line by line, putting in relevant information from dataframe and handle necessary formatting.
        # For each row, we have 1 page of Word doc, 22 lines. After the last paragraph, next letter starts in a new page.

        line0 = letter.add_paragraph(dt.datetime.now().strftime('%m/%d/%y'))

        line1 = letter.add_paragraph('RE: ' +
                                     str(entry['Address']) + ', ' + str(entry['City']) + ', ' +
                                     str(entry['State']) + ', ' + str(entry['Zip']) + ', ')

        line2 = letter.add_paragraph(str(entry['Owner 1 First Name']) + ' ' + str(entry['Owner 1 Last Name']))

        line3 = letter.add_paragraph(str(entry['Mailing Address']) + ' ' + str(entry['Mailing Unit #']))

        line4 = letter.add_paragraph(str(entry['Mailing City']) + ', ' + str(entry['Mailing State']) + ', ' + str(entry['Mailing Zip']))

        line5 = letter.add_paragraph('Dear ' + str(entry['Owner 1 First Name']) + str(entry['Owner 1 Last Name']) + ',')

        line6 = letter.add_paragraph('My name is ')
        line6.add_run('John Peter, ').bold = True
        line6.add_run('I am interested in buying your apartment building at ' + str(entry['Address']))

        line7 = letter.add_paragraph('I am an experienced local investor and currently own several investment properties. '
                                     'I have a down payment and I’m looking to buy larger properties …just like yours. '
                                     'I’m also pre-approved with XYZ Bank. I got your name through a public records search. I am ')
        line7.add_run('NOT ').bold = True
        line7.add_run('a realtor.')

        line8 = letter.add_paragraph()
        line8.add_run('When I buy a property').bold = True

        line9 = letter.add_paragraph('You won’t need to fool around listing your property '
                                     'with a realtor and all the Headaches that go along with it.', style='List Bullet')
        line9.paragraph_format.left_indent = Inches(0.5)
        line10 = letter.add_paragraph('I will buy the property As-Is. '
                                      'You won’t need to fix anything before you sell.', style='List Bullet')
        line10.paragraph_format.left_indent = Inches(0.5)
        line11 = letter.add_paragraph('You won’t have to pay realtor commissions. You save thousand$!', style='List Bullet')
        line11.paragraph_format.left_indent = Inches(0.5)
        line12 = letter.add_paragraph('I can complete the sale quickly. '
                                      'You don’t have to worry if your buyer will actually be able to close.', style='List Bullet')
        line12.paragraph_format.left_indent = Inches(0.5)
        line13 = letter.add_paragraph('Avoid endless negotiations and an uncertain sales price; '
                                      'you will know exactly what you will sell it for.', style='List Bullet')
        line13.paragraph_format.left_indent = Inches(0.5)

        line14 = letter.add_paragraph('My goal is to make this very easy for you. I will remove all the stress!')

        line15 = letter.add_paragraph()
        line15.add_run('Please Call or text me at ').bold = True
        phone_number = line15.add_run('123-456-789 ')
        phone_number.bold = True ; phone_number.underline = True
        line15.add_run('right away!').bold = True

        line16 = letter.add_paragraph()
        line16.add_run('Or email me at ').bold = True
        email_address = line16.add_run('john@peter.com ')
        email_address.bold = True ; email_address.underline = True
        line16.add_run('for a quick response.').bold = True

        line17 = letter.add_paragraph('I am also looking at other properties and I can only buy one or two right now!')

        line18 = letter.add_paragraph('Call me right away to have yours considered!')

        line19 = letter.add_paragraph('Thanks, ')

        line20 = letter.add_paragraph()
        line20.add_run('John Peter').bold = True

        line21 = letter.add_paragraph()
        phone_number2 = line21.add_run('123-456-789')
        phone_number2.bold = True ; phone_number2.underline = True

        line22 = letter.add_paragraph('P.S. If I am unable to answer when you call, please leave a voice mail and I will get back ASAP!')

    letter.save('output.docx')